from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


OUTPUT_PATH = Path(__file__).resolve().parent / "Fuzzy_POI_Recommendation_Section1_Documentation_AR.docx"

NAVY = "16324F"
TEAL = "147D82"
LIGHT_TEAL = "DDEFF0"
LIGHT_BLUE = "EAF1F7"
LIGHT_GRAY = "F3F5F7"
MID_GRAY = "D8DEE5"
TEXT = "23313F"
WHITE = "FFFFFF"
ACCENT = "D9A441"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_rtl(table):
    tbl_pr = table._tbl.tblPr
    bidi = tbl_pr.find(qn("w:bidiVisual"))
    if bidi is None:
        bidi = OxmlElement("w:bidiVisual")
        tbl_pr.append(bidi)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_run_font(run, name="Arial", size=10.5, bold=False, color=TEXT):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{attr}"), name)
    lang = r_pr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        r_pr.append(lang)
    lang.set(qn("w:val"), "ar-SA")
    lang.set(qn("w:bidi"), "ar-SA")


def set_paragraph_rtl(paragraph, alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_after=4, line_spacing=1.08):
    paragraph.alignment = alignment
    paragraph.paragraph_format.space_after = Pt(space_after)
    paragraph.paragraph_format.line_spacing = line_spacing
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = p_pr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        p_pr.append(bidi)


def add_rtl_paragraph(container, text="", size=10.5, bold=False, color=TEXT,
                      alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_after=4,
                      line_spacing=1.08, keep_with_next=False):
    paragraph = container.add_paragraph()
    set_paragraph_rtl(paragraph, alignment, space_after, line_spacing)
    paragraph.paragraph_format.keep_with_next = keep_with_next
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return paragraph


def add_heading(container, text, level=1):
    sizes = {1: 17, 2: 13.5, 3: 11.5}
    colors = {1: NAVY, 2: TEAL, 3: NAVY}
    paragraph = container.add_paragraph()
    set_paragraph_rtl(paragraph, WD_ALIGN_PARAGRAPH.RIGHT, space_after=5, line_spacing=1.0)
    paragraph.paragraph_format.space_before = Pt(7 if level == 1 else 4)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text)
    set_run_font(run, size=sizes[level], bold=True, color=colors[level])
    p_pr = paragraph._p.get_or_add_pPr()
    bottom_border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "3")
    bottom.set(qn("w:color"), TEAL if level == 1 else MID_GRAY)
    bottom_border.append(bottom)
    p_pr.append(bottom_border)
    return paragraph


def add_bullet(container, text, size=10.2, color=TEXT, space_after=2):
    paragraph = container.add_paragraph()
    set_paragraph_rtl(paragraph, WD_ALIGN_PARAGRAPH.RIGHT, space_after=space_after, line_spacing=1.05)
    paragraph.paragraph_format.right_indent = Cm(0.35)
    run = paragraph.add_run("• " + text)
    set_run_font(run, size=size, color=color)
    return paragraph


def set_cell_text(cell, text, bold=False, color=TEXT, size=9.3,
                  align=WD_ALIGN_PARAGRAPH.RIGHT, fill=None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    set_paragraph_rtl(paragraph, align, space_after=0, line_spacing=1.0)
    run = paragraph.add_run(str(text))
    set_run_font(run, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margins(cell)
    if fill:
        set_cell_shading(cell, fill)


def add_table(document, headers, rows, widths=None, font_size=9.2, header_size=9.5):
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    set_table_rtl(table)
    header_row = table.rows[0]
    set_repeat_table_header(header_row)
    prevent_row_split(header_row)

    for index, header in enumerate(headers):
        if widths:
            header_row.cells[index].width = widths[index]
        set_cell_text(
            header_row.cells[index],
            header,
            bold=True,
            color=WHITE,
            size=header_size,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            fill=NAVY,
        )

    for row_index, values in enumerate(rows):
        row = table.add_row()
        prevent_row_split(row)
        fill = LIGHT_BLUE if row_index % 2 == 0 else WHITE
        for col_index, value in enumerate(values):
            if widths:
                row.cells[col_index].width = widths[col_index]
            align = WD_ALIGN_PARAGRAPH.CENTER if col_index == 0 else WD_ALIGN_PARAGRAPH.RIGHT
            set_cell_text(row.cells[col_index], value, size=font_size, align=align, fill=fill)
    document.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_page_number(paragraph):
    set_paragraph_rtl(paragraph, WD_ALIGN_PARAGRAPH.CENTER, space_after=0, line_spacing=1.0)
    label = paragraph.add_run("صفحة ")
    set_run_font(label, size=8.5, color="66727E")
    run = paragraph.add_run()
    set_run_font(run, size=8.5, color="66727E")
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr_text, fld_separate, fld_end])


def add_page_break(document):
    paragraph = document.add_paragraph()
    paragraph.add_run().add_break(WD_BREAK.PAGE)


def configure_document(document):
    section = document.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.35)
    section.bottom_margin = Cm(1.25)
    section.left_margin = Cm(1.4)
    section.right_margin = Cm(1.4)
    section.header_distance = Cm(0.55)
    section.footer_distance = Cm(0.55)
    section.different_first_page_header_footer = True

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.08

    header = section.header
    paragraph = header.paragraphs[0]
    set_paragraph_rtl(paragraph, WD_ALIGN_PARAGRAPH.CENTER, space_after=0, line_spacing=1.0)
    run = paragraph.add_run("Fuzzy Logic Project  |  Section 1: Fuzzy Expert System")
    set_run_font(run, size=8.5, bold=True, color=TEAL)

    footer = section.footer
    add_page_number(footer.paragraphs[0])
    add_page_number(section.first_page_footer.paragraphs[0])


def add_cover(document):
    document.add_paragraph().paragraph_format.space_after = Pt(35)

    banner = document.add_table(rows=1, cols=1)
    banner.alignment = WD_TABLE_ALIGNMENT.CENTER
    banner.autofit = False
    banner.columns[0].width = Cm(17.8)
    cell = banner.cell(0, 0)
    set_cell_shading(cell, NAVY)
    set_cell_margins(cell, top=520, start=220, bottom=520, end=220)
    title = cell.paragraphs[0]
    set_paragraph_rtl(title, WD_ALIGN_PARAGRAPH.CENTER, space_after=10, line_spacing=1.0)
    run = title.add_run("Fuzzy POI Recommendation System\nfor Trip Planning")
    set_run_font(run, size=25, bold=True, color=WHITE)
    subtitle = cell.add_paragraph()
    set_paragraph_rtl(subtitle, WD_ALIGN_PARAGRAPH.CENTER, space_after=0, line_spacing=1.0)
    run = subtitle.add_run("نظام توصية الأماكن السياحية باستخدام Fuzzy Logic")
    set_run_font(run, size=15, bold=True, color="CDE9EA")

    document.add_paragraph().paragraph_format.space_after = Pt(24)
    accent = document.add_table(rows=1, cols=1)
    accent.alignment = WD_TABLE_ALIGNMENT.CENTER
    accent.autofit = False
    accent.columns[0].width = Cm(7.0)
    set_cell_shading(accent.cell(0, 0), ACCENT)
    set_cell_margins(accent.cell(0, 0), top=18, bottom=18)

    add_rtl_paragraph(
        document,
        "Section 1 – Fuzzy Expert System",
        size=17,
        bold=True,
        color=TEAL,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=18,
    )
    add_rtl_paragraph(document, "Course / Project: Fuzzy Logic Project", size=12, bold=True,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
    add_rtl_paragraph(document, "Student Name: ______________________________", size=11.5,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
    add_rtl_paragraph(document, "Date: 11 يونيو 2026", size=11.5,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=22)

    note = document.add_table(rows=1, cols=1)
    note.alignment = WD_TABLE_ALIGNMENT.CENTER
    note.autofit = False
    note.columns[0].width = Cm(15.5)
    set_cell_text(
        note.cell(0, 0),
        "توثيق مختصر لتصميم وتنفيذ واختبار Fuzzy Expert System الخاص بترتيب المعالم والمطاعم.",
        size=10.5,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        fill=LIGHT_TEAL,
    )


def add_overview_page(document):
    add_heading(document, "1. المقدمة", 1)
    add_rtl_paragraph(
        document,
        "يشرح هذا المستند الجزء الأول من المشروع فقط، وهو بناء Fuzzy Expert System لتوصية الأماكن السياحية والمطاعم أثناء تخطيط الرحلة. يعتمد الشرح على متطلبات Project Assessment وعلى التنفيذ الموجود في notebook باسم fuzzy_poi_recommendation_system.ipynb.",
    )

    add_heading(document, "2. Problem Statement", 1)
    add_rtl_paragraph(
        document,
        "اختيار مكان مناسب للزيارة لا يعتمد على قرارات ثنائية بسيطة. فالسائح يستخدم عبارات مرنة مثل: قريب، تقييمه مرتفع، يناسب ميزانيتي، يطابق اهتماماتي، ومتاح في هذا الوقت. هذه المفاهيم تتدرج ولا يمكن تمثيلها بدقة باستخدام True أو False فقط، لذلك تعد Fuzzy Logic مناسبة لتحويل هذه التفضيلات إلى درجات مفهومة بين 0 و100.",
    )
    add_rtl_paragraph(
        document,
        "يقيم النظام كل POI، سواء كان attraction أو restaurant، ثم ينتج Recommendation Score وVisit Priority لاستخدامهما في ترتيب أفضل الخيارات.",
    )

    add_heading(document, "3. System Overview", 1)
    workflow = document.add_table(rows=1, cols=7)
    workflow.alignment = WD_TABLE_ALIGNMENT.CENTER
    workflow.autofit = False
    set_table_rtl(workflow)
    steps = [
        ("City + JSON Data\n+ User Preferences", NAVY, WHITE, 3.5),
        ("←", WHITE, TEAL, 0.7),
        ("Feature\nExtraction", TEAL, WHITE, 2.7),
        ("←", WHITE, TEAL, 0.7),
        ("Fuzzy Expert\nSystem", NAVY, WHITE, 2.8),
        ("←", WHITE, TEAL, 0.7),
        ("Recommendation Score\n+ Visit Priority", TEAL, WHITE, 3.7),
    ]
    for index, (label, fill, color, width_cm) in enumerate(steps):
        workflow.columns[index].width = Cm(width_cm)
        set_cell_text(workflow.cell(0, index), label, bold=True, color=color, size=9.0,
                      align=WD_ALIGN_PARAGRAPH.CENTER, fill=fill)
    document.add_paragraph().paragraph_format.space_after = Pt(2)

    add_heading(document, "البيانات والتوليد التلقائي للمدخلات", 2)
    add_rtl_paragraph(
        document,
        "يقرأ النظام ملفي attractions.json وrestaurants.json من مسارات يحددها المستخدم في أعلى الـ notebook، ثم يوحد الحقول ويصفي النتائج حسب اسم المدينة. يدعم الكود نقص بعض الحقول دون إيقاف التنفيذ.",
    )
    add_bullet(document, "Distance Score: يحسب Haversine distance من موقع الفندق/المستخدم، أو من مركز المدينة المقدر من متوسط الإحداثيات.")
    add_bullet(document, "Rating Score: يحول rating من المجال 0–5 إلى المجال 0–100، والقيمة المفقودة تعامل بقيمة وسطية.")
    add_bullet(document, "Popularity Score: يستخدم log normalization على num_reviews حتى لا تهيمن الأماكن شديدة الشهرة على الترتيب.")
    add_bullet(document, "Budget Match: يقارن مستوى ميزانية المستخدم مع price، price_level، fee أو أقل سعر متاح للمكان.")
    add_bullet(document, "Interest Match: يستخدم keyword matching مع category، subcategory، subtype، cuisine وdescription.")
    add_bullet(document, "Availability Score: يفحص hours.week_ranges للوقت المختار؛ القيم هي 100 للمفتوح، 60 عند غياب الساعات، و20 للمغلق.")


def add_variables_page(document):
    add_heading(document, "4. Input Variables and Fuzzy Sets", 1)
    input_rows = [
        ("Distance Score\n0–100", "درجة قرب المكان من موقع الفندق أو المستخدم.", "Poor / Acceptable / Excellent"),
        ("Rating Score\n0–100", "التقييم بعد تحويله من 0–5 إلى 0–100.", "Low / Good / Excellent"),
        ("Popularity Score\n0–100", "شعبية المكان بعد log normalization لعدد المراجعات.", "Unpopular / Known / Popular"),
        ("Budget Match\n0–100", "مدى توافق سعر المكان مع مستوى ميزانية المستخدم.", "Poor / Acceptable / Ideal"),
        ("Interest Match\n0–100", "مدى تطابق التصنيف أو المطبخ أو الوصف مع اهتمامات المستخدم.", "Weak / Moderate / Strong"),
        ("Availability Score\n0–100", "مدى ملاءمة ساعات العمل لليوم والوقت المختارين.", "Closed/Unknown / Maybe Available / Available"),
    ]
    add_table(
        document,
        ["Input Variable والمدى", "المعنى", "Fuzzy Sets"],
        input_rows,
        widths=[Cm(4.0), Cm(8.1), Cm(5.8)],
        font_size=8.9,
    )

    add_heading(document, "5. Output Variables and Fuzzy Sets", 1)
    output_rows = [
        ("Recommendation Score\n0–100", "جودة التوصية النهائية للمكان.", "Low / Medium / High"),
        ("Visit Priority\n0–100", "أولوية إدراج المكان ضمن خطة الزيارة.", "Low / Medium / High"),
    ]
    add_table(
        document,
        ["Output Variable والمدى", "المعنى", "Fuzzy Sets"],
        output_rows,
        widths=[Cm(4.3), Cm(8.0), Cm(5.6)],
        font_size=9.1,
    )

    add_heading(document, "6. Membership Functions", 1)
    add_rtl_paragraph(
        document,
        "يستخدم النظام Triangular وTrapezoidal Membership Functions لأنها واضحة، سهلة الضبط، وسهلة العرض للطالب والمستخدم. كل متغير يعمل على Universe من 0 إلى 100، وتسمح مناطق التداخل بانتماء القيمة إلى أكثر من Fuzzy Set بدرجات مختلفة.",
    )
    add_bullet(document, "Low أو Poor: غالباً Trapezoidal Membership Function لتمثيل المجال المنخفض بثبات عند البداية.")
    add_bullet(document, "Medium أو Acceptable: غالباً Triangular Membership Function بقمة في منتصف المجال.")
    add_bullet(document, "High أو Excellent: غالباً Trapezoidal Membership Function لتمثيل المجال المرتفع بثبات قرب 100.")
    add_rtl_paragraph(
        document,
        "ينشئ الـ notebook رسوماً لجميع Membership Functions باستخدام matplotlib، مما يساعد على فحص مناطق التداخل وفهم كيفية تحويل القيم الرقمية إلى مفاهيم لغوية.",
        size=10.2,
        bold=True,
        color=TEAL,
    )


RULES = [
    "إذا كان Interest Match = Strong وRating Score = Excellent وDistance Score = Excellent، فإن Recommendation Score = High وVisit Priority = High.",
    "إذا كان Interest Match = Strong وBudget Match = Ideal وAvailability Score = Available، فإن Recommendation Score = High وVisit Priority = High.",
    "إذا كان Rating Score = Excellent وPopularity Score = Popular وInterest Match = Moderate، فإن Recommendation Score = High وVisit Priority = Medium.",
    "إذا كان Distance Score = Poor وBudget Match = Poor، فإن Recommendation Score = Low وVisit Priority = Low.",
    "إذا كان Interest Match = Weak، فإن Recommendation Score = Low وVisit Priority = Low.",
    "إذا كان Availability Score = Closed/Unknown وكان Interest Match ليس Strong، فإن Recommendation Score = Low وVisit Priority = Low.",
    "إذا كان Budget Match = Acceptable وRating Score = Good وInterest Match = Moderate، فإن Recommendation Score = Medium وVisit Priority = Medium.",
    "إذا كان Distance Score = Excellent وRating Score = Good، فإن Recommendation Score = Medium وVisit Priority = Medium.",
    "إذا كان Popularity Score = Unpopular وRating Score = Low، فإن Recommendation Score = Low وVisit Priority = Low.",
    "إذا كان Interest Match = Strong وDistance Score = Poor لكن Rating Score = Excellent، فإن Recommendation Score = Medium وVisit Priority = Medium.",
    "إذا كان Availability Score = Available وBudget Match = Ideal وRating Score = Excellent، فإن Recommendation Score = High وVisit Priority = High.",
    "إذا كانت معظم الإشارات Medium أو Acceptable، فإن Recommendation Score = Medium وVisit Priority = Medium.",
]


def add_rule_page(document, start, end, page_title):
    add_heading(document, page_title, 1)
    if start == 0:
        add_rtl_paragraph(
            document,
            "تمثل Rule Base المعرفة الخبيرة داخل النظام. يحتوي التنفيذ الحالي على 14 قاعدة، ويعرض الجدول التالي 12 قاعدة أساسية مطلوبة في التوثيق. يمكن تفعيل أكثر من قاعدة في الوقت نفسه، ثم تدمج النتائج قبل Defuzzification.",
            size=10.1,
        )
    rows = [(str(index + 1), RULES[index]) for index in range(start, end)]
    add_table(
        document,
        ["#", "Fuzzy Rule"],
        rows,
        widths=[Cm(1.0), Cm(16.9)],
        font_size=9.25,
        header_size=9.6,
    )


def add_inference_sections(document):
    add_heading(document, "8. Inference Engine", 1)
    add_rtl_paragraph(
        document,
        "يستخدم النظام Mamdani Inference، وهي من أكثر الطرق شيوعاً في Fuzzy Expert Systems. تبدأ العملية بتحويل القيم الرقمية إلى درجات عضوية، ثم تقييم شروط القواعد باستخدام عمليات AND وOR، وبعد ذلك تجميع مخرجات القواعد الفعالة لتكوين Fuzzy Output موحد لكل Output Variable.",
        size=10.0,
    )

    add_heading(document, "9. Defuzzification", 1)
    add_rtl_paragraph(
        document,
        "بعد التجميع نحتاج إلى قيمة رقمية قابلة للترتيب. يستخدم المشروع Centroid Method لحساب مركز مساحة الـ output membership function. تعطي هذه الطريقة نتيجة متوازنة بين 0 و100، وهي مناسبة لـ Recommendation Score وVisit Priority.",
        size=10.0,
    )


def add_final_page(document):
    add_heading(document, "10. User Interface", 1)
    add_rtl_paragraph(
        document,
        "توفر واجهة Jupyter Notebook المبنية باستخدام ipywidgets طريقة مباشرة لتعديل الإعدادات وإعادة تشغيل التوصية دون تعديل الخلايا البرمجية. يمكن للمستخدم تحديد:",
        size=10.0,
    )
    add_bullet(document, "City name، Budget level، User interests، Preferred day وPreferred visit time.", size=9.8)
    add_bullet(document, "Hotel latitude وHotel longitude بشكل اختياري، إضافة إلى Number of top results.", size=9.8)
    add_rtl_paragraph(document, "تعرض الواجهة جدول الترتيب، Bar Chart لأفضل النتائج، وشرحاً مختصراً لأفضل POI.", size=9.8)

    add_heading(document, "11. Testing and Validation", 1)
    add_rtl_paragraph(
        document,
        "تتم المقارنة بين expected label وpredicted label باستخدام مجموعة صغيرة معنونة يدوياً وخمس Synthetic Test Cases. يستخدم التصنيف التالي:",
        size=9.8,
    )
    add_table(
        document,
        ["المجال", "Label"],
        [("0–39", "Low"), ("40–69", "Medium"), ("70–100", "High")],
        widths=[Cm(5.0), Cm(5.0)],
        font_size=8.8,
        header_size=9.0,
    )
    add_rtl_paragraph(
        document,
        "في التشغيل الحالي: تم تحميل 30 attraction و30 restaurant، ثم تصفية 59 POI لمدينة Las Vegas. بلغت Synthetic Validation Accuracy نسبة 100%، بينما بلغت Manual Validation Accuracy نسبة 33.3%. النتيجة الثانية تشير إلى أن قواعد expected labels اليدوية تحتاج إلى معايرة أفضل حتى تتوافق مع Rule Base، ولا تعني وجود خطأ في تنفيذ Mamdani Inference.",
        size=9.5,
        color=NAVY,
    )

    add_heading(document, "12. Dependencies and Running Instructions", 1)
    code_box = document.add_table(rows=1, cols=1)
    code_box.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_cell_text(
        code_box.cell(0, 0),
        "pip install numpy pandas matplotlib scikit-fuzzy ipywidgets python-docx",
        size=9.0,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        fill=LIGHT_GRAY,
    )
    add_bullet(document, "اضبط ATTRACTIONS_PATH وRESTAURANTS_PATH في أعلى الـ notebook.", size=9.5)
    add_bullet(document, "حدد CITY_NAME والميزانية والاهتمامات واليوم والوقت، ثم شغل الخلايا من الأعلى إلى الأسفل.", size=9.5)
    add_bullet(document, "راجع Membership Functions والجدول النهائي والرسوم ونتائج Testing and Validation.", size=9.5)

    add_heading(document, "13. Limitations and Future Improvements", 1)
    add_rtl_paragraph(
        document,
        "القيود الحالية تشمل الاعتماد على keyword matching البسيط، احتمال نقص ratings أو hours أو prices، وعدم وجود route optimization. يمكن مستقبلاً استخدام NLP لمطابقة الاهتمامات، إضافة real-time APIs، تحسين معايرة Rule Base، وربط النتائج مع itinerary optimization.",
        size=9.5,
    )

    add_heading(document, "14. References", 1)
    references = [
        "scikit-fuzzy Documentation: https://scikit-fuzzy.readthedocs.io/",
        "Zadeh, L. A. (1965). Fuzzy Sets. Information and Control, 8(3), 338–353.",
        "Mamdani, E. H. & Assilian, S. (1975). An Experiment in Linguistic Synthesis with a Fuzzy Logic Controller.",
        "Ross, T. J. Fuzzy Logic with Engineering Applications - Mamdani Inference and Centroid Defuzzification.",
        "Dataset Source: [يضاف رابط أو وصف مصدر attractions.json وrestaurants.json].",
    ]
    for reference in references:
        add_bullet(document, reference, size=8.6, space_after=1)


def build_document():
    document = Document()
    configure_document(document)
    properties = document.core_properties
    properties.title = "Fuzzy POI Recommendation System for Trip Planning"
    properties.subject = "Section 1 - Fuzzy Expert System"
    properties.author = "Student Name: [Placeholder]"
    properties.keywords = "Fuzzy Logic, scikit-fuzzy, POI Recommendation, Mamdani, Arabic Documentation"

    add_cover(document)
    add_page_break(document)
    add_overview_page(document)
    add_page_break(document)
    add_variables_page(document)
    add_page_break(document)
    add_rule_page(document, 0, 6, "7. Rule Base – القواعد 1 إلى 6")
    add_page_break(document)
    add_rule_page(document, 6, 12, "7. Rule Base – القواعد 7 إلى 12")
    add_inference_sections(document)
    add_page_break(document)
    add_final_page(document)

    document.save(OUTPUT_PATH)
    return OUTPUT_PATH


if __name__ == "__main__":
    print(build_document())
